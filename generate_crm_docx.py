import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def set_cell_background(cell, hex_color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def create_hubzzoo_crm_business_plan_docx(filename):
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    navy = RGBColor(15, 23, 42)
    blue = RGBColor(37, 99, 235)
    dark_gray = RGBColor(51, 65, 85)
    emerald = RGBColor(16, 185, 129)
    
    # Title
    p_title = doc.add_paragraph()
    r_title = p_title.add_run("FORRETNINGSPLAN")
    r_title.bold = True
    r_title.font.size = Pt(26)
    r_title.font.color.rgb = navy
    p_title.paragraph_format.space_after = Pt(2)

    p_sub = doc.add_paragraph()
    r_sub = p_sub.add_run("Hubzzoo (hubzzoo.ai.studio) — AI Executive CRM & Sales Platform")
    r_sub.bold = True
    r_sub.font.size = Pt(15)
    r_sub.font.color.rgb = blue
    p_sub.paragraph_format.space_after = Pt(14)
    
    # Metadata Table
    meta_table = doc.add_table(rows=4, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        ("Dokumentversjon:", "1.0 (Offisiell lanseringsplan)"),
        ("Selskap:", "Hubzzoo CRM / aiappsy"),
        ("Forfatter / Eier:", "Pål Alexander Juritzen"),
        ("Dato:", "August 2026")
    ]
    for i, (k, v) in enumerate(meta_data):
        row = meta_table.rows[i]
        c1, c2 = row.cells[0], row.cells[1]
        c1.text = k
        c2.text = v
        c1.paragraphs[0].runs[0].bold = True
        c1.paragraphs[0].runs[0].font.size = Pt(9.5)
        c2.paragraphs[0].runs[0].font.size = Pt(9.5)
        set_cell_background(c1, "F1F5F9")
        set_cell_background(c2, "F8FAFC")
        set_cell_margins(c1, top=60, bottom=60, left=100, right=100)
        set_cell_margins(c2, top=60, bottom=60, left=100, right=100)
        c1.width = Inches(2.0)
        c2.width = Inches(4.5)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    
    def add_h1(text):
        h = doc.add_paragraph()
        r = h.add_run(text)
        r.bold = True
        r.font.size = Pt(16)
        r.font.color.rgb = navy
        h.paragraph_format.space_before = Pt(16)
        h.paragraph_format.space_after = Pt(6)
        return h

    def add_h2(text):
        h = doc.add_paragraph()
        r = h.add_run(text)
        r.bold = True
        r.font.size = Pt(12.5)
        r.font.color.rgb = blue
        h.paragraph_format.space_before = Pt(10)
        h.paragraph_format.space_after = Pt(4)
        return h

    def add_body(text, bold=False):
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.bold = bold
        r.font.size = Pt(10.5)
        r.font.color.rgb = dark_gray
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        return p

    def add_bullet(text, bold_prefix=""):
        p = doc.add_paragraph(style='List Bullet')
        if bold_prefix:
            r1 = p.add_run(bold_prefix)
            r1.bold = True
            r1.font.size = Pt(10)
            r1.font.color.rgb = navy
        r2 = p.add_run(text)
        r2.font.size = Pt(10)
        r2.font.color.rgb = dark_gray
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15

    # 1. Executive Summary
    add_h1("1. Sammendrag (Executive Summary)")
    add_body("Hubzzoo (hubzzoo.ai.studio) er et komplett, neste-generasjons AI-drevet CRM- og salgsautomatiseringssystem utviklet for moderne bedrifter, konsulenter og salgsteam.")
    add_body("Systemet kombinerer Google Gemini AI-assistanse, visuell Nexus-leadfangst, automatisert e-postutsending og direkte Stripe-betalingslenker for å automatisere hele salgsprosessen fra første kontakt til signert kontrakt og betaling.")

    add_h2("1.1 Visjon & Hovedmålsetting")
    add_bullet(" Nå 500 betalende B2B-kunder og $35 000 i månedlig gjentakende inntekt (MRR).", "År 1:")
    add_bullet(" Skalere til 2 500 aktive bedriftskunder og $175 000 i MRR med whitelabel-støtte.", "År 2:")
    add_bullet(" Nå 8 000 bedriftskunder med en årlig omsetning (ARR) på over $4,8 millioner og 92 % driftsmargin.", "År 3:")

    # 2. Produkt & Kjernefunksjoner
    add_h1("2. Produkt & Kjernefunksjonalitet")
    add_bullet(" Innebygd Gemini-drevet assistent som analyserer innkommende leads, undersøker selskapsinformasjon og skriver skreddersydde tilbud.", "1. Executive AI Assistant:")
    add_bullet(" Visuelt verktøy for å bygge konverterende landingssider og skjemaer som sender leads direkte inn i CRM-et.", "2. Nexus Visual Lead Capture:")
    add_bullet(" Automatisk generering og utsending av Stripe-betalingslenker for umiddelbart oppgjør.", "3. Integrert Stripe-fakturering:")
    add_bullet(" Automatisk oppfølging av leads via SMTP/Nodemailer basert på kundeadferd.", "4. Smart E-postautomatisering:")
    add_bullet(" Skalerbar, sanntids skydatabase bygget for enterprise-sikkerhet og null forsinkelse.", "5. Firebase & Cloud Firestore:")

    # 3. Inntektsmodell
    add_h1("3. Inntektsmodell & Pristiers")
    
    t_rev = doc.add_table(rows=4, cols=3)
    t_rev.alignment = WD_TABLE_ALIGNMENT.CENTER
    rev_headers = ["Abonnementsnivå", "Funksjoner", "Pris (MRR)"]
    for j, h_text in enumerate(rev_headers):
        c = t_rev.rows[0].cells[j]
        c.text = h_text
        c.paragraphs[0].runs[0].bold = True
        c.paragraphs[0].runs[0].font.color.rgb = navy
        c.paragraphs[0].runs[0].font.size = Pt(9.5)
        set_cell_background(c, "E2E8F0")
        set_cell_margins(c, top=80, bottom=80, left=120, right=120)
        
    rev_rows = [
        ("Starter CRM", "Opptil 500 kontakter, grunnleggende pipeline og e-postintegrasjon.", "$29 / måned"),
        ("Executive PRO", "Ubegrensede kontakter, full Gemini AI-assistent og Stripe-betalinger.", "$79 / måned"),
        ("Enterprise / Whitelabel", "Skreddersydde AI-agenter, eget domene, dedikert onboarding og team-lisenser.", "$199 / måned")
    ]
    for i, (k1, k2, k3) in enumerate(rev_rows):
        row = t_rev.rows[i+1]
        for idx, val in enumerate([k1, k2, k3]):
            cell = row.cells[idx]
            cell.text = val
            cell.paragraphs[0].runs[0].font.size = Pt(9)
            if idx == 0: cell.paragraphs[0].runs[0].bold = True
            if idx == 2: 
                cell.paragraphs[0].runs[0].bold = True
                cell.paragraphs[0].runs[0].font.color.rgb = blue
            set_cell_background(cell, "FFFFFF" if i%2==0 else "F8FAFC")
            set_cell_margins(cell, top=60, bottom=60, left=100, right=100)

    # 4. Finansielle Projeksjoner (Tabell)
    add_h1("4. Finansielle Projeksjoner (3 År)")
    
    t_fin = doc.add_table(rows=7, cols=4)
    t_fin.alignment = WD_TABLE_ALIGNMENT.CENTER
    fin_headers = ["Nøkkeltall / Metrikk", "År 1", "År 2", "År 3"]
    for j, h_text in enumerate(fin_headers):
        c = t_fin.rows[0].cells[j]
        c.text = h_text
        c.paragraphs[0].runs[0].bold = True
        c.paragraphs[0].runs[0].font.color.rgb = navy
        c.paragraphs[0].runs[0].font.size = Pt(9.5)
        set_cell_background(c, "E2E8F0")
        set_cell_margins(c, top=80, bottom=80, left=120, right=120)

    fin_rows = [
        ("Betalende Bedriftskunder", "500", "2 500", "8 000"),
        ("Gjennomsnittlig Månedspris (ARPU)", "$68", "$74", "$82"),
        ("Årlige Abonnementsinntekter (ARR)", "$408 000", "$2 220 000", "$7 872 000"),
        ("AppSumo / Lifetime Deal Salg", "$65 000", "$0 (Utfaset)", "$0"),
        ("Totale Bruttoinntekter", "$473 000", "$2 220 000", "$7 872 000"),
        ("Netto Driftsresultat (EBITDA)", "$435 000", "$2 060 000", "$7 320 000")
    ]
    for i, row_data in enumerate(fin_rows):
        row = t_fin.rows[i+1]
        for idx, val in enumerate(row_data):
            cell = row.cells[idx]
            cell.text = val
            cell.paragraphs[0].runs[0].font.size = Pt(9)
            if idx == 0: cell.paragraphs[0].runs[0].bold = True
            if idx > 0 and i in [4, 5]: 
                cell.paragraphs[0].runs[0].bold = True
                cell.paragraphs[0].runs[0].font.color.rgb = blue if i==4 else emerald
            set_cell_background(cell, "FFFFFF" if i%2==0 else "F8FAFC")
            set_cell_margins(cell, top=60, bottom=60, left=100, right=100)

    doc.save(filename)
    print(f"Created: {filename}")

def create_hubzzoo_crm_sales_listing_docx(filename):
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    navy = RGBColor(15, 23, 42)
    blue = RGBColor(37, 99, 235)
    emerald = RGBColor(16, 185, 129)
    dark_gray = RGBColor(51, 65, 85)

    # Title
    p_title = doc.add_paragraph()
    r_title = p_title.add_run("SALGSPROSPEKT & ANNONSE")
    r_title.bold = True
    r_title.font.size = Pt(24)
    r_title.font.color.rgb = navy
    p_title.paragraph_format.space_after = Pt(2)

    p_sub = doc.add_paragraph()
    r_sub = p_sub.add_run("Hubzzoo (hubzzoo.ai.studio) — AI Executive CRM [Nøkkelferdig SaaS]")
    r_sub.bold = True
    r_sub.font.size = Pt(14)
    r_sub.font.color.rgb = blue
    p_sub.paragraph_format.space_after = Pt(14)

    def add_h1(text):
        h = doc.add_paragraph()
        r = h.add_run(text)
        r.bold = True
        r.font.size = Pt(15)
        r.font.color.rgb = navy
        h.paragraph_format.space_before = Pt(14)
        h.paragraph_format.space_after = Pt(4)
        return h

    def add_body(text, bold=False):
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.bold = bold
        r.font.size = Pt(10.5)
        r.font.color.rgb = dark_gray
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        return p

    def add_bullet(text, bold_prefix=""):
        p = doc.add_paragraph(style='List Bullet')
        if bold_prefix:
            r1 = p.add_run(bold_prefix)
            r1.bold = True
            r1.font.size = Pt(10)
            r1.font.color.rgb = navy
        r2 = p.add_run(text)
        r2.font.size = Pt(10)
        r2.font.color.rgb = dark_gray
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15

    add_body("Er du på jakt etter å overta en fullt ferdigutviklet, skalerbar B2B AI-programvarevirksomhet med over 92 % fortjenestemargin?")
    add_body("Nå har du muligheten til å sikre deg Hubzzoo CRM (hubzzoo.ai.studio) – et komplett, moderne CRM- og salgsautomatiseringssystem bygget på React 19, Express, Firebase Firestore, Stripe og Google Gemini AI.")

    # Pricing Box Table
    add_h1("Finansielle Betingelser & Prissetting")
    
    t_price = doc.add_table(rows=3, cols=3)
    t_price.alignment = WD_TABLE_ALIGNMENT.CENTER
    p_headers = ["Element", "Anbefalt Pris", "Hva Kjøper Får"]
    for j, h_text in enumerate(p_headers):
        c = t_price.rows[0].cells[j]
        c.text = h_text
        c.paragraphs[0].runs[0].bold = True
        c.paragraphs[0].runs[0].font.color.rgb = navy
        c.paragraphs[0].runs[0].font.size = Pt(9.5)
        set_cell_background(c, "E2E8F0")
        set_cell_margins(c, top=80, bottom=80, left=120, right=120)

    p_rows = [
        ("1. Kjøpesum for Virksomheten (Engangsbeløp)", "kr 129 000,- NOK\n($12 900 USD)", "100 % full eiendomsrett til hele Hubzzoo CRM-systemet, AI Executive Assistant, Nexus Editor, Firebase blueprint, Stripe-integrasjon og 3-års forretningsplan."),
        ("2. Månedlig Drifts- & Supportavtale (SLA)", "kr 3 990,- / mnd\n($390 USD/mnd)", "Full drift og overvåking på Google Cloud Run, Firebase databasedrift, Gemini AI API-kvoter, teknisk feilretting og prioritert support.")
    ]
    for i, row_data in enumerate(p_rows):
        row = t_price.rows[i+1]
        for idx, val in enumerate(row_data):
            cell = row.cells[idx]
            cell.text = val
            cell.paragraphs[0].runs[0].font.size = Pt(9)
            if idx == 0: cell.paragraphs[0].runs[0].bold = True
            if idx == 1: 
                cell.paragraphs[0].runs[0].bold = True
                cell.paragraphs[0].runs[0].font.color.rgb = blue if i==0 else emerald
            set_cell_background(cell, "FFFFFF" if i%2==0 else "F8FAFC")
            set_cell_margins(cell, top=60, bottom=60, left=100, right=100)

    add_h1("Hva Følger Med i Kjøpet (100 % Full Eiendomsrett):")
    add_bullet(" Intelligent AI-assistent drevet av Google Gemini som analyserer leads, gjør selskapsresearch og genererer tilbud på sekunder.", "1. Executive AI Assistant:")
    add_bullet(" Visuell landingsside-bygger som samler leads automatisk inn i salgspipelinen.", "2. Nexus Visual Lead Capture:")
    add_bullet(" Direkte opprettelse og utsending av Stripe-betalingslenker fra CRM-et.", "3. Stripe Betalingsintegrasjon:")
    add_bullet(" Automatiserte e-postsekvenser via Nodemailer/SMTP basert på kundestatus.", "4. Smart E-postoppfølging:")
    add_bullet(" Komplett 3-års forretningsplan, Go-To-Market plan og finansiell modell.", "5. Dokumentasjon & Forretningsplan:")

    add_h1("Kontakt & Overdragelse:")
    add_bullet(" Signering av kjøpekontrakt og overdragelsesavtale.", "1. Avtale:")
    add_bullet(" Overføring av GitHub-kildekode, Firebase-prosjekt, tilganger og domener.", "2. Overføring:")
    add_bullet(" 1-til-1 digital onboarding (2 timer) for full teknisk og operasjonell innføring.", "3. Onboarding:")
    add_body("\nKontakt for demo og oversendelse av fullt prospekt:\nE-post: paljuritzen@gmail.com | GitHub: github.com/aiappsy/aistudio-crm-app", bold=True)

    doc.save(filename)
    print(f"Created: {filename}")

if __name__ == "__main__":
    out_dir = r"C:\Users\paul\.gemini\antigravity\scratch\aistudio-crm-app"
    bp_path = os.path.join(out_dir, "Forretningsplan_Hubzzoo_CRM.docx")
    sales_path = os.path.join(out_dir, "Salgsannonse_Hubzzoo_CRM.docx")
    
    create_hubzzoo_crm_business_plan_docx(bp_path)
    create_hubzzoo_crm_sales_listing_docx(sales_path)
